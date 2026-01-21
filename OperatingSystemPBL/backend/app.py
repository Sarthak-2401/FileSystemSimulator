# backend/app.py
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
import math
import os
import sqlite3
import hashlib
import gzip
import shutil
from datetime import datetime
from math import ceil
from flask_cors import cross_origin
# Additional libraries for file conversion
import whisper
import pdfplumber
from docx import Document

from sentence_transformers import SentenceTransformer, util
import torch

doc = Document()

# model = SentenceTransformer("all-MiniLM-L6-v2")
model = SentenceTransformer('all-MiniLM-L6-v2')  # your embedding model
whisper_model = whisper.load_model("base")       # for MP3 transcription


app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": ["http://localhost:5173", "http://127.0.0.1:5173"]}})


# @app.after_request
# def add_cors_headers(response):
#     response.headers.add("Access-Control-Allow-Origin", "*")
#     response.headers.add("Access-Control-Allow-Headers", "Content-Type,Authorization")
#     response.headers.add("Access-Control-Allow-Methods", "GET,POST,PUT,DELETE,OPTIONS")
#     return response


# ---------- CONFIG ----------
UPLOAD_DIR = "uploads"
DB_FILE = "database.db"
TOTAL_BLOCKS = 1000
BLOCK_SIZE_KB = 4  # 4KB
JUNK_EXTENSIONS = ['.tmp', '.log', '.bak', '.cache']
os.makedirs(UPLOAD_DIR, exist_ok=True)



# ---------- DB HELPERS ----------
def get_conn():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_conn()
    c = conn.cursor()
    # files: metadata + hash, compression
    c.execute("""
    CREATE TABLE IF NOT EXISTS files (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        filename TEXT,
        stored_filename TEXT,
        size_kb REAL,
        original_size_kb REAL,
        uploaded_at TEXT,
        allocation_type TEXT,
        is_compressed INTEGER DEFAULT 0,
        sha256 TEXT
    )""")
    # blocks: block index -> file
    c.execute("""
    CREATE TABLE IF NOT EXISTS blocks (
        block_index INTEGER PRIMARY KEY,
        file_id INTEGER,
        next_block INTEGER,
        FOREIGN KEY(file_id) REFERENCES files(id)
    )""")
    # logs / journal
    c.execute("""
    CREATE TABLE IF NOT EXISTS logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        action TEXT,
        timestamp TEXT
    )""")
    # ai_recommendations
    c.execute("""
    CREATE TABLE IF NOT EXISTS ai_recommendations (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        file_id INTEGER,
        suggestion TEXT,
        confidence REAL,
        created_at TEXT
    )""")
    conn.commit()
    conn.close()

init_db()

# ---------- UTIL ----------
def add_log(action):
    conn = get_conn()
    c = conn.cursor()
    c.execute("INSERT INTO logs (action, timestamp) VALUES (?, ?)", (action, datetime.utcnow().isoformat()))
    conn.commit()
    conn.close()

def compute_sha256(path, chunk_size=8192):
    h = hashlib.sha256()
    with open(path, "rb") as f:
        while True:
            chunk = f.read(chunk_size)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()

def get_free_blocks():
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT block_index FROM blocks WHERE file_id IS NULL OR file_id = 0 ORDER BY block_index")
    rows = c.fetchall()
    conn.close()
    return [r['block_index'] for r in rows]

def ensure_blocks_table_populated():
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT COUNT(*) as cnt FROM blocks")
    cnt = c.fetchone()['cnt']
    if cnt < TOTAL_BLOCKS:
        # fill missing blocks
        existing = set()
        c.execute("SELECT block_index FROM blocks")
        for r in c.fetchall():
            existing.add(r['block_index'])
        for i in range(TOTAL_BLOCKS):
            if i not in existing:
                c.execute("INSERT OR REPLACE INTO blocks (block_index, file_id, next_block) VALUES (?, ?, ?)", (i, None, None))
        conn.commit()
    conn.close()

ensure_blocks_table_populated()

def fragmentation_percent():
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT file_id, COUNT(*) as contiguous_count FROM blocks WHERE file_id IS NOT NULL GROUP BY file_id")
    used_block_count = 0
    fragments = 0
    c.execute("SELECT block_index, file_id FROM blocks ORDER BY block_index")
    rows = c.fetchall()
    prev_file = None
    for r in rows:
        fid = r['file_id']
        if fid is not None:
            used_block_count += 1
            if prev_file is not None and prev_file != fid:
                fragments += 1
        prev_file = fid
    conn.close()
    if TOTAL_BLOCKS == 0:
        return 0.0
    return (fragments / TOTAL_BLOCKS) * 100.0

# ---------- ALLOCATION HELPERS ----------
def find_contiguous(num_blocks):
    conn = get_conn()
    c = conn.cursor()
    # scan blocks table for contiguous sequence of free blocks
    c.execute("SELECT block_index, file_id FROM blocks ORDER BY block_index")
    rows = c.fetchall()
    start = None
    length = 0
    for r in rows:
        if r['file_id'] is None:
            if start is None:
                start = r['block_index']
                length = 1
            else:
                length += 1
            if length >= num_blocks:
                conn.close()
                return start
        else:
            start = None
            length = 0
    conn.close()
    return -1

def find_free_blocks_any(num_blocks):
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT block_index FROM blocks WHERE file_id IS NULL ORDER BY block_index LIMIT ?", (num_blocks,))
    rows = [r['block_index'] for r in c.fetchall()]
    conn.close()
    return rows if len(rows) == num_blocks else []

def occupy_blocks(file_id, blocks_list):
    conn = get_conn()
    c = conn.cursor()
    # mark blocks assigned and set next pointers
    for i, b in enumerate(blocks_list):
        next_b = blocks_list[i+1] if i+1 < len(blocks_list) else None
        c.execute("UPDATE blocks SET file_id = ?, next_block = ? WHERE block_index = ?", (file_id, next_b, b))
    conn.commit()
    conn.close()




@app.route("/delete/<int:file_id>", methods=["DELETE", "OPTIONS"])
@cross_origin(origin='http://localhost:5173', methods=['DELETE', 'OPTIONS'])
def delete_file(file_id):
    if request.method == "OPTIONS":
        return '', 200  # Handle preflight request for CORS

    try:
        conn = get_conn()
        c = conn.cursor()

        # Get stored file name before deleting DB entry
        c.execute("SELECT stored_filename FROM files WHERE id = ?", (file_id,))
        row = c.fetchone()
        stored_filename = row["stored_filename"] if row else None

        # Delete file record and free its blocks
        c.execute("DELETE FROM files WHERE id = ?", (file_id,))
        c.execute("UPDATE blocks SET file_id = NULL, next_block = NULL WHERE file_id = ?", (file_id,))
        conn.commit()
        conn.close()

        # Remove actual file from uploads directory
        if stored_filename:
            file_path = os.path.join(UPLOAD_DIR, stored_filename)
            if os.path.exists(file_path):
                os.remove(file_path)

        add_log(f"Deleted file with id {file_id}")
        print(f"✅ Deleted file {file_id} successfully")

        return jsonify({"message": f"File {file_id} deleted successfully"}), 200

    except Exception as e:
        print(f"❌ Error deleting file {file_id}: {e}")
        return jsonify({"error": str(e)}), 500



@app.route("/logs", methods=["GET"])
def get_logs():
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM logs ORDER BY id DESC")
    rows = c.fetchall()
    conn.close()

    logs = []
    for r in rows:
        logs.append({
            "id": r["id"],
            "action": r["action"],
            "timestamp": r["timestamp"]
        })
    return jsonify(logs), 200


@app.route("/files", methods=["GET"])
def get_files():
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM files ORDER BY id DESC")
    rows = c.fetchall()
    conn.close()

    files = []
    for r in rows:
        files.append({
            "id": r["id"],
            "filename": r["filename"],
            "size_kb": r["size_kb"],
            "allocation_type": r["allocation_type"],
            "uploaded_at": r["uploaded_at"]
        })
    return jsonify(files), 200


@app.route("/init", methods=["GET"])
def reset_filesystem():
    # delete existing DB
    if os.path.exists(DB_FILE):
        os.remove(DB_FILE)

    # reinitialize tables
    init_db()

    # repopulate block table so all blocks show as free
    ensure_blocks_table_populated()

    # journal entry
    add_log("System reset: filesystem reinitialized.")

    return jsonify({"message": "File system reset successfully", "ok": True}), 200



# @app.route("/upload", methods=["POST"])
# def upload():
#     print("DEBUG: request.files =", request.files)
#     print("DEBUG: request.form =", request.form)
#     ensure_blocks_table_populated()

#     if "file" not in request.files:
#         return jsonify({"error": "No file part"}), 400

#     file = request.files["file"]
#     allocation_type = request.form.get("allocation_type", "contiguous")

#     if file.filename == "":
#         return jsonify({"error": "No selected file"}), 400

#     filename = secure_filename(file.filename)
#     stored_name = f"{int(datetime.utcnow().timestamp())}_{filename}"
#     file_path = os.path.join(UPLOAD_DIR, stored_name)
#     file.save(file_path)

#     size_kb = os.path.getsize(file_path) / 1024.0
#     original_size_kb = size_kb
#     sha = compute_sha256(file_path)

#     # number of blocks needed
#     num_blocks = max(1, math.ceil(size_kb / BLOCK_SIZE_KB))

#     # -------- SELECT ALLOCATION STRATEGY --------
#     if allocation_type == "contiguous":
#         start = find_contiguous(num_blocks)
#         if start == -1:
#             return jsonify({"error": "Not enough contiguous space"}), 400
#         blocks_list = list(range(start, start + num_blocks))

#     elif allocation_type == "linked":
#         blocks_list = find_free_blocks_any(num_blocks)
#         if not blocks_list:
#             return jsonify({"error": "Not enough free blocks"}), 400

#     elif allocation_type == "indexed":
#         blocks_list = find_free_blocks_any(num_blocks)
#         if not blocks_list:
#             return jsonify({"error": "Not enough free blocks"}), 400

#     else:
#         return jsonify({"error": "Invalid allocation type"}), 400

#     # insert file record
#     conn = get_conn()
#     c = conn.cursor()
#     c.execute("""
#         INSERT INTO files (filename, stored_filename, size_kb, original_size_kb, uploaded_at, allocation_type, sha256)
#         VALUES (?, ?, ?, ?, ?, ?, ?)
#     """, (filename, stored_name, size_kb, original_size_kb, datetime.utcnow().isoformat(), allocation_type, sha))
#     file_id = c.lastrowid
#     conn.commit()
#     conn.close()

#     # occupy blocks
#     occupy_blocks(file_id, blocks_list)

#     add_log(f"Uploaded file '{filename}' using {allocation_type} allocation.")
#     return jsonify({"message": "File uploaded", "file_id": file_id, "blocks": blocks_list}), 200

@app.route("/upload", methods=["POST"])
def upload():
    print("DEBUG: request.files =", request.files)
    print("DEBUG: request.form =", request.form)
    ensure_blocks_table_populated()

    # ✅ Use .get() to safely fetch the file
    file = request.files.get("file")
    allocation_type = request.form.get("allocation_type", "contiguous")

    if not file:
        return jsonify({"error": "No file uploaded"}), 400
    if file.filename == "":
        return jsonify({"error": "No selected file"}), 400

    # Secure filename & save
    filename = secure_filename(file.filename)
    stored_name = f"{int(datetime.utcnow().timestamp())}_{filename}"
    file_path = os.path.join(UPLOAD_DIR, stored_name)
    file.save(file_path)

    size_kb = os.path.getsize(file_path) / 1024.0
    original_size_kb = size_kb
    sha = compute_sha256(file_path)

    # number of blocks needed
    num_blocks = max(1, math.ceil(size_kb / BLOCK_SIZE_KB))

    # -------- SELECT ALLOCATION STRATEGY --------
    blocks_list = []
    if allocation_type == "contiguous":
        start = find_contiguous(num_blocks)
        if start == -1:
            return jsonify({"error": "Not enough contiguous space"}), 400
        blocks_list = list(range(start, start + num_blocks))

    elif allocation_type in ["linked", "indexed"]:
        blocks_list = find_free_blocks_any(num_blocks)
        if not blocks_list:
            return jsonify({"error": "Not enough free blocks"}), 400

    else:
        return jsonify({"error": "Invalid allocation type"}), 400

    # insert file record
    conn = get_conn()
    c = conn.cursor()
    c.execute("""
        INSERT INTO files (filename, stored_filename, size_kb, original_size_kb, uploaded_at, allocation_type, sha256)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (filename, stored_name, size_kb, original_size_kb, datetime.utcnow().isoformat(), allocation_type, sha))
    file_id = c.lastrowid
    conn.commit()
    conn.close()

    # occupy blocks
    occupy_blocks(file_id, blocks_list)

    add_log(f"Uploaded file '{filename}' using {allocation_type} allocation.")
    return jsonify({"message": "File uploaded", "file_id": file_id, "blocks": blocks_list}), 200



# @app.route('/optimize', methods=['POST'])
# def optimize():
#     # Load stored files from DB (file paths)
#     conn = get_conn()
#     cur = conn.cursor()
#     cur.execute("SELECT id, name, path FROM files")
#     files = cur.fetchall()

#     if len(files) < 2:
#         return {"message": "Not enough files to optimize"}, 200

#     # Read file contents
#     contents = []
#     file_ids = []
#     for file_id, name, path in files:
#         try:
#             with open(path, "rb") as f:
#                 content = f.read().decode(errors="ignore")
#             contents.append(content)
#             file_ids.append((file_id, name))
#         except:
#             continue

#     # Generate embeddings
#     embeddings = model.encode(contents, convert_to_tensor=True)

#     # Compute similarity matrix
#     similarity_matrix = util.cos_sim(embeddings, embeddings)

#     suggestions = []
#     threshold = 0.80  # 80% similarity = potential duplicate

#     for i in range(len(files)):
#         for j in range(i + 1, len(files)):
#             score = float(similarity_matrix[i][j])
#             if score >= threshold:
#                 suggestions.append({
#                     "file1": file_ids[i][1],
#                     "file2": file_ids[j][1],
#                     "similarity": round(score * 100, 2)
#                 })

#     return {"duplicates": suggestions}, 200


# def mp3_to_text(path):
#     try:
#         result = whisper_model.transcribe(path)
#         return result['text']
#     except Exception as e:
#         print(f"MP3 transcription failed for {path}: {e}")
#         return ""

# def pdf_to_text(path):
#     text = ""
#     try:
#         with pdfplumber.open(path) as pdf:
#             for page in pdf.pages:
#                 text += page.extract_text() + "\n"
#     except Exception as e:
#         print(f"PDF reading failed for {path}: {e}")
#     return text

# def docx_to_text(path):
#     text = ""
#     try:
#         doc = docx.Document(path)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     except Exception as e:
#         print(f"DOCX reading failed for {path}: {e}")
#     return text

@app.route('/optimize', methods=['POST'])
def optimize():
    conn = get_conn()
    cur = conn.cursor()

    # Fetch all files
    cur.execute("SELECT id, filename, stored_filename FROM files")
    files = cur.fetchall()

    if len(files) < 2:
        return {"message": "Not enough files to optimize"}, 200

    contents = []
    file_ids = []

    # Helper functions
    def mp3_to_text(path):
        result = whisper_model.transcribe(path)
        return result['text']

    def pdf_to_text(path):
        import pdfplumber
        content = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                content += page.extract_text() or ""
        return content

    def docx_to_text(path):
        import docx
        doc = docx.Document(path)
        return "\n".join([p.text for p in doc.paragraphs])

    # Process files
       # Process files
    for file in files:
        file_id = file['id']
        filename = file['filename']
        path = os.path.join(UPLOAD_DIR, file['stored_filename'])

        try:
            ext = os.path.splitext(filename)[1].lower()
            content = ""

            if ext == ".txt":
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()

            elif ext == ".mp3":
                content = mp3_to_text(path)

            elif ext == ".pdf":
                content = pdf_to_text(path)

            elif ext == ".docx":
                content = docx_to_text(path)

            # ✅ New: support for images and subtitles
            elif ext in [".png", ".jpg", ".jpeg"]:
                content = f"[IMAGE FILE: {filename}]"

            elif ext in [".srt", ".vtt"]:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()

            else:
                print(f"Unsupported file type: {filename}")
                continue

            if content.strip() == "":
                continue

            contents.append(content)
            file_ids.append((file_id, filename))

        except Exception as e:
            print(f"Failed to process {filename}: {e}")
            continue


    if not contents:
        return {"message": "No valid files to optimize"}, 200

    # Generate embeddings and similarity
    embeddings = model.encode(contents, convert_to_tensor=True)
    similarity_matrix = util.cos_sim(embeddings, embeddings)

    suggestions = []
    threshold = 0.80  # 80% similarity = potential duplicate

    for i in range(len(file_ids)):
        for j in range(i + 1, len(file_ids)):
            score = float(similarity_matrix[i][j])
            if score >= threshold:
                suggestions.append({
                    "file1": file_ids[i][1],
                    "file2": file_ids[j][1],
                    "similarity": round(score * 100, 2)
                })

    return {"duplicates": suggestions}, 200


@app.route("/blocks", methods=["GET"])
def get_blocks():
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT block_index, file_id, next_block FROM blocks ORDER BY block_index")
    rows = c.fetchall()
    conn.close()

    block_list = []
    for r in rows:
        block_list.append({
            "block_index": r["block_index"],
            "file_id": r["file_id"],
            "next_block": r["next_block"]
        })

    return jsonify({"blocks": block_list}), 200



# ---------- DEFRAg ----------
def defragment():
    """
    Simple defragment — collect files in upload order and reassign contiguous blocks.
    """
    ensure_blocks_table_populated()
    conn = get_conn()
    c = conn.cursor()
    # clear all block allocations
    c.execute("UPDATE blocks SET file_id = NULL, next_block = NULL")
    conn.commit()

    # get files ordered by id (upload order)
    c.execute("SELECT id FROM files ORDER BY id")
    files = [r['id'] for r in c.fetchall()]

    current = 0
    for fid in files:
        # find how many blocks were used previously (we saved blocks_count maybe; compute from size)
        c.execute("SELECT size_kb FROM files WHERE id = ?", (fid,))
        row = c.fetchone()
        if not row:
            continue
        size_kb = row['size_kb']
        num_blocks = ceil(size_kb / BLOCK_SIZE_KB) or 1
        blocks = list(range(current, current + num_blocks))
        occupy_blocks(fid, blocks)
        current += num_blocks
    conn.commit()
    conn.close()
    add_log("Defragmentation complete")

@app.route("/fragmentation", methods=["GET"])
def get_fragmentation():
    return jsonify({"fragmentation": fragmentation_percent()})

@app.route("/defragment", methods=["POST"])
def defragment_endpoint():
    defragment()
    return jsonify({"message": "Defragmentation complete"}), 200


# ---------- simple recommendations endpoint (AI-ish) ----------
@app.route("/recommendations", methods=["GET"])
def recommendations():
    """
    Return heuristic suggestions per file:
     - compress if > threshold
     - delete if junk
     - mark duplicate groups
    """
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT id, filename, size_kb, is_compressed, sha256 FROM files")
    rows = c.fetchall()
    # build duplicates map
    sha_map = {}
    for r in rows:
        sha_map.setdefault(r['sha256'], []).append(r['id'])

    recs = []
    for r in rows:
        rid = r['id']
        fname = r['filename']
        ext = os.path.splitext(fname)[1].lower()
        size_kb = r['size_kb']
        if ext in JUNK_EXTENSIONS:
            recs.append({"file_id": rid, "suggestion": "delete (junk)", "confidence": 0.95})
        elif not r['is_compressed'] and size_kb > 200:
            recs.append({"file_id": rid, "suggestion": "compress", "confidence": 0.85})
        elif len(sha_map.get(r['sha256'], [])) > 1:
            recs.append({"file_id": rid, "suggestion": "duplicate - consider delete", "confidence": 0.9})
    conn.close()
    return jsonify(recs)

# ---------- init endpoint ----------
@app.route("/init", methods=["POST"])
def reset_all():
    """
    Warning: destroys data — useful for dev/testing
    """
    ensure_blocks_table_populated()
    conn = get_conn()
    c = conn.cursor()
    c.execute("DELETE FROM files")
    c.execute("UPDATE blocks SET file_id = NULL, next_block = NULL")
    c.execute("DELETE FROM logs")
    conn.commit()
    conn.close()
    # delete files on disk
    for fname in os.listdir(UPLOAD_DIR):
        try:
            os.remove(os.path.join(UPLOAD_DIR, fname))
        except Exception:
            pass
    add_log("System initialized (reset)")
    return jsonify({"ok": True})

# ---------- run ----------
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
